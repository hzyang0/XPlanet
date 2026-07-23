from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_interview_handbook_docx as base


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "XPlanet-初版后端面试学习手册.md"
OUTPUT = ROOT / "docs" / "XPlanet-初版后端面试学习手册.docx"


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(70)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("XPLANET · INITIAL BACKEND EDITION")
    base.set_run_font(run, 10.5, True, base.BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("XPlanet 初版后端面试学习手册")
    base.set_run_font(run, 25, True, base.INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("专业技能详解 · 项目请求链路 · 高频追问 · 故障分析 · 复习计划")
    base.set_run_font(run, 12.5, color=base.DARK_BLUE)

    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.paragraph_format.left_indent = Inches(0.42)
    callout.paragraph_format.right_indent = Inches(0.42)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(36)
    base.add_inline(
        callout,
        "以 Git 初版 `9bcaeb5` 为事实基线：只讲简历对应实现，后续优化仅作为演进方案",
        size=10.5,
        color="344054",
    )
    base.shade(callout._p, base.CALLOUT_FILL)

    for text in (
        "适用方向：Java 后端开发 / 秋招项目深挖 / 后端基础复习",
        "建议用法：先画流程，再学原理，随后遮住答案模拟追问，最后亲手做故障实验",
        "版本：2026-07-23",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(text)
        base.set_run_font(run, 9.5, color=base.MUTED)

    doc.add_page_break()


def configure_header_footer(doc: Document) -> None:
    base.configure_page(doc)
    section = doc.sections[0]
    # compact_reference_guide preset: Letter, 1-inch margins, 0.492-inch
    # header/footer distance, yielding the required 6.5-inch / 9360-DXA body.
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("XPlanet 初版  |  Java 后端面试学习手册")
    base.set_run_font(run, 9, color=base.MUTED)


def restart_numbered_lists(doc: Document) -> None:
    """Give each contiguous List Number block a real numbering instance starting at 1."""
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]
    style_num_pr = style._element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    base_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    next_num_id = max(
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ) + 1

    active_num_id: int | None = None
    for paragraph in doc.paragraphs:
        if paragraph.style is None or paragraph.style.name != "List Number":
            active_num_id = None
            continue
        if active_num_id is None:
            active_num_id = next_num_id
            next_num_id += 1
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(active_num_id))
            abstract = OxmlElement("w:abstractNumId")
            abstract.set(qn("w:val"), str(abstract_num_id))
            num.append(abstract)
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), "0")
            start = OxmlElement("w:startOverride")
            start.set(qn("w:val"), "1")
            override.append(start)
            num.append(override)
            numbering.append(num)

        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        for child in list(num_pr):
            num_pr.remove(child)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(active_num_id))
        num_pr.extend((ilvl, num_id))


def remove_forced_chapter_breaks(doc: Document) -> None:
    """Let chapters flow naturally; keep_with_next still prevents orphan headings."""
    for paragraph in doc.paragraphs:
        if paragraph.style is not None and paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.page_break_before = False


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    base.configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "XPlanet 初版后端面试学习手册"
    doc.core_properties.subject = "Java 后端专业技能、XPlanet 初版项目实现与面试问答"
    doc.core_properties.author = "XPlanet Project"
    doc.core_properties.keywords = (
        "XPlanet, Java, Spring Boot, MySQL, Redis, RocketMQ, Caffeine, "
        "JUC, JVM, 缓存一致性, 点赞幂等"
    )

    add_cover(doc)
    base.add_toc(doc, headings)
    base.add_body(doc, lines)
    restart_numbered_lists(doc)
    remove_forced_chapter_breaks(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
