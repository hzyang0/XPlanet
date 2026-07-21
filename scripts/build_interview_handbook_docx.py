from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "INTERVIEW-GUIDE.md"
OUTPUT = ROOT / "docs" / "XPlanet-Research-秋招面试八股手册.docx"

INK = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "F3F4F6"
GRID = "B8C4D1"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, color=None, name=BODY_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, latin, cjk, size, color=None, bold=None):
    style.font.name = latin
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def shade(element, fill):
    props = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)


def add_inline(paragraph, text, size=None, color=None):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size or 9.5, color=DARK_BLUE, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_FONT, CJK_FONT, 11, "202124")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, BODY_FONT, CJK_FONT, 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, BODY_FONT, CJK_FONT, 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, BODY_FONT, CJK_FONT, 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, BODY_FONT, CJK_FONT, 11, "202124")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("XPlanet Research  |  秋招面试八股手册")
    set_run_font(run, 9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(86)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("XPLANET RESEARCH")
    set_run_font(run, 11, True, BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("秋招面试八股手册")
    set_run_font(run, 28, True, INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(32)
    run = subtitle.add_run("围绕真实项目实现的 Java 后端 + AI Agent 面试参考")
    set_run_font(run, 13.5, color=DARK_BLUE)

    guide = doc.add_paragraph()
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide.paragraph_format.left_indent = Inches(0.55)
    guide.paragraph_format.right_indent = Inches(0.55)
    guide.paragraph_format.space_after = Pt(40)
    add_inline(
        guide,
        "先给结论 · 再讲项目实现 · 解释设计取舍 · 最后说明边界与演进",
        size=11,
        color="344054",
    )
    shade(guide._p, CALLOUT_FILL)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("适用方向：Java 后端 / Agent 应用开发 / Java + AI 复合岗位")
    set_run_font(run, 10, color=MUTED)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("基于当前可运行仓库基线  ·  2026-07-21")
    set_run_font(run, 9.5, color=MUTED)

    doc.add_page_break()


def add_toc(doc, headings):
    h = doc.add_paragraph("目录", style="Heading 1")
    h.paragraph_format.page_break_before = False
    h.paragraph_format.space_after = Pt(12)
    for index, heading in enumerate(headings, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, re.sub(r"^\d+\.\s*", "", heading), size=10.5, color=INK)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def choose_widths(rows):
    cols = len(rows[0])
    headers = rows[0]
    if cols == 2:
        if headers[0] in {"错误说法", "键", "问题"}:
            return [3000, 6360]
        return [4200, 5160]
    if cols == 3:
        return [1500, 2800, 5060]
    if cols == 4:
        return [1600, 2500, 2600, 2660]
    base = 9360 // cols
    return [base] * (cols - 1) + [9360 - base * (cols - 1)]


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, choose_widths(rows))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, text, size=9.2 if cols >= 3 else 9.6, color=INK if row_index == 0 else "202124")
            if row_index == 0:
                shade(cell._tc, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif cols >= 3 and col_index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc, code, language):
    if language == "mermaid":
        code = (
            "浏览器工作台 → Gateway → User / Article / Interaction / AI\n"
            "AI → MySQL（Task / Run / Outbox）→ RocketMQ → Consumer → Python Agent\n"
            "Agent → internal_search / Web Search / Fetch → Evidence / Claim / Citation\n"
            "人工审核 → 幂等发布文章 → 进入后续站内检索"
        )
    p = doc.add_paragraph()
    shade(p._p, CODE_FILL)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, 9, color="344054", name="Consolas")


def add_body(doc, lines):
    start = next(index for index, line in enumerate(lines) if line.startswith("## "))
    index = start
    in_code = False
    code_language = ""
    code_lines = []
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("~~~"):
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines), code_language)
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            p.paragraph_format.page_break_before = index != start
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:], style="Heading 3")
        elif re.match(r"^- \[[ xX]\] ", stripped):
            checked = stripped[3].lower() == "x"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            add_inline(p, ("☑ " if checked else "☐ ") + stripped[6:], size=10.5)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade(p._p, CALLOUT_FILL)
            add_inline(p, stripped.lstrip("> "), size=10.5, color="344054")
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        index += 1


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "XPlanet Research 秋招面试八股手册"
    doc.core_properties.subject = "Java 后端与 AI Agent 秋招面试参考"
    doc.core_properties.author = "XPlanet Project"
    doc.core_properties.keywords = "XPlanet, Java, Agent, LangGraph, Outbox, RocketMQ, Redis"

    add_cover(doc)
    add_toc(doc, headings)
    add_body(doc, lines)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
